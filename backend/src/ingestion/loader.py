"""loader.py — Đọc & parse Sổ tay sinh viên (DOCX) giữ cấu trúc Phần/Chương/Mục/Điều.

Đặc điểm tài liệu thực tế (khảo sát 2026-08, file ICTU 2024-2025):
- ~2349 paragraph + 26 bảng XEN KẼ trong body → phải duyệt body.iterchildren()
  thay vì doc.paragraphs, nếu không sẽ mất toàn bộ bảng.
- Heading style KHÔNG thống nhất (Heading 1-4, 'Heading #2' lẫn 'Normal')
  → nhận diện cấu trúc bằng REGEX + style-map, không dựa vào style thuần.
- "Chương" có 2 tầng: CHƯƠNG I/II (IN HOA) là chương ngoài của Phần;
  "Chương I...V" (thường, trong một Mục) là chương CON của riêng quy chế đó.
- Số Điều đánh lại từ 1 trong mỗi văn bản con (mục A, B, C, D...)
  → metadata phải kèm muc để phân biệt "Điều 1" của quy chế nào.
- Dòng Mục lục có 3 dạng: tab+số trang, dấu chấm+số trang, số dính liền chữ
  → lọc theo style 'toc*' + regex, không dùng trạng thái vùng TOC.

Metadata mỗi Section:
    phan, chuong, chuong_con, muc, trich, dieu, ten_dieu,
    so_trang (nội suy theo offset ký tự), nguon
"""
from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# --- Regex nhận diện cấu trúc pháp lý ---
RE_PHAN = re.compile(r"^PHẦN\s+([IVX]{1,4}|\d+)\b[.:]?\s*(.*)", re.I)
RE_CHUONG_OUT = re.compile(r"^CHƯƠNG\s+([IVX]{1,4}|\d+)\b[.:]?\s*(.*)")      # chương ngoài (IN HOA)
RE_CHUONG_IN = re.compile(r"^Chương\s+([IVX]{1,4}|\d+)\b[.:]?\s*(.*)")        # chương con trong 1 quy chế
RE_MUC = re.compile(
    r"^([A-D])\.\s+(Quy định|Quy chế|Quy trình|Đối với|Trích|Hướng dẫn|Quyết định)\b(.{0,200})"
)
RE_TRICH = re.compile(r"^\d*\.\s*(Trích|TRÍCH)\s+(Thông tư|Quyết định|Nghị định)\s+số\s+([\w/\-]+)")
RE_DIEU = re.compile(r"^Điều\s+(\d+)\s*[.:\)]\s*(.*)")

# --- Regex lọc dòng MỤC LỤC ---
RE_TOC_TITLE = re.compile(r"^\s*(MỤC LỤC|PHỤ LỤC)\s*$", re.I)
RE_TOC_DOTS = re.compile(r"[.…\.]{2,}\s*\d{1,3}\s*$")          # "...155"
RE_TOC_KEYWORD = re.compile(r"^\s*(PHẦN|CHƯƠNG|Chương|[A-D]\.)\b.*\d{1,3}\s*$")  # số trang dính liền

# Dòng trang trí: toàn gạch ngang / dấu chấm
RE_DECOR = re.compile(r"^[\-–—=.*…\s]+$")

FALLBACK_CHARS_PER_PAGE = 1800   # DOCX nhiều bảng/ảnh: ~1800 ký tự/trang in

# app.xml đôi khi lưu số trang SAI (VD: Word lưu "3" cho tài liệu 160 trang).
# Chỉ tin giá trị app.xml nếu hàm ý số ký tự/trang nằm trong khoảng hợp lý.
MIN_CHARS_PER_PAGE = 300
MAX_CHARS_PER_PAGE = 4000


@dataclass
class Section:
    """Một đoạn văn bản logic (thường = 1 Điều, hoặc tiêu đề/giới thiệu)."""
    text: str
    phan: str = ""
    chuong: str = ""
    chuong_con: str = ""   # chương con bên trong 1 quy chế (VD: "Chương II" của Quy chế rèn luyện)
    muc: str = ""          # mục A/B/C/D — tên quy chế/quy định con
    trich: str = ""        # văn bản được trích dẫn (Thông tư, Quyết định...)
    dieu: str = ""         # "Điều 5"
    ten_dieu: str = ""     # "Quyền của người học"
    so_trang: int = 0      # nội suy, KHÔNG phải số trang thật của bản in
    nguon: str = ""
    _offset: int = 0       # offset ký tự, dùng để nội suy số trang


@dataclass
class ParsedDocument:
    file_path: str
    total_chars: int = 0
    tong_so_trang: int = 0   # đọc từ docProps/app.xml (0 nếu không có)
    n_sections: int = 0
    sections: list = field(default_factory=list)


def _para_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.findall(".//" + qn("w:t"))).strip()


def _table_text(tbl_el) -> str:
    """Render bảng thành text dạng 'hàng | cột' để giữ cấu trúc cho LLM đọc."""
    lines = []
    for row in tbl_el.findall(qn("w:tr")):
        cells = []
        for tc in row.findall(qn("w:tc")):
            cell = " ".join(
                "".join(t.text or "" for t in p.findall(".//" + qn("w:t")))
                for p in tc.findall(qn("w:p"))
            ).strip()
            cells.append(cell)
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_pages_from_app_xml(path: Path) -> int:
    """Đọc số trang Word lưu trong docProps/app.xml (giá trị tham khảo)."""
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("docProps/app.xml").decode("utf-8", "ignore")
        m = re.search(r"<Pages>(\d+)</Pages>", xml)
        return int(m.group(1)) if m else 0
    except Exception:
        return 0


def _is_toc_line(text: str, style: str) -> bool:
    """Nhận diện dòng mục lục: style 'toc*' HOẶC kết thúc bằng số trang."""
    if style.lower().startswith("toc"):
        return True
    return bool(RE_TOC_DOTS.search(text) or RE_TOC_KEYWORD.match(text))


def load_docx(path: str | Path) -> ParsedDocument:
    """Parse DOCX thành ParsedDocument (các Section kèm metadata phân cấp).

    Thuật toán: state machine duyệt tuần tự body (paragraph + bảng xen kẽ).
    Marker cấu trúc (Phần/Chương/Mục/Trích/Điều) nhận diện bằng regex;
    tiêu đề đa dòng được gộp qua cơ chế pending_title.
    """
    path = Path(path)
    doc = Document(str(path))
    body = doc.element.body

    # Style-map: element w:p -> tên style (style trong file này không đáng tin
    # để xác định cấu trúc, nhưng hữu ích để lọc dòng mục lục)
    style_of = {p._p: (p.style.name if p.style else "") for p in doc.paragraphs}

    result = ParsedDocument(file_path=str(path))
    result.tong_so_trang = _read_pages_from_app_xml(path)

    st = {k: "" for k in ("phan", "chuong", "chuong_con", "muc", "trich", "dieu", "ten_dieu")}
    pending_title, pending_kind = "", ""
    char_count = 0
    buf: list[str] = []
    buf_offset = 0

    def flush() -> None:
        nonlocal buf, buf_offset
        text = "\n".join(buf).strip()
        if text:
            result.sections.append(Section(text=text, **st, so_trang=0,
                                           nguon=path.name, _offset=buf_offset))
        buf, buf_offset = [], char_count

    def commit_pending() -> None:
        nonlocal pending_title, pending_kind
        if pending_kind == "phan" and pending_title:
            st["phan"] = f"{st['phan']} – {pending_title}"
        elif pending_kind == "chuong" and pending_title:
            st["chuong"] = f"{st['chuong']}: {pending_title}"
        elif pending_kind == "chuong_con" and pending_title:
            st["chuong_con"] = f"{st['chuong_con']}: {pending_title}"
        pending_title, pending_kind = "", ""

    def append_content(text: str) -> None:
        nonlocal char_count
        if not buf:
            buf_offset = char_count
        char_count += len(text) + 1
        buf.append(text)

    for el in body.iterchildren():
        if el.tag == qn("w:p"):
            text = _para_text(el)
            if not text:
                continue
            style = style_of.get(el, "")
            if RE_TOC_TITLE.match(text) or _is_toc_line(text, style):
                continue
            if RE_DECOR.match(text):
                continue

            # ---------- Marker: ĐIỀU ----------
            m = RE_DIEU.match(text)
            if m:
                flush()
                commit_pending()
                st["dieu"] = f"Điều {m.group(1)}"
                st["ten_dieu"] = m.group(2).strip()
                append_content(text)
                continue

            # ---------- Marker: PHẦN (IN HOA) ----------
            m = RE_PHAN.match(text)
            if m:
                flush()
                commit_pending()
                st["phan"] = f"Phần {m.group(1)}"
                st.update(chuong="", chuong_con="", muc="", trich="", dieu="", ten_dieu="")
                rest = m.group(2).strip()
                pending_title, pending_kind = (rest, "phan") if rest else ("", "")
                continue

            # ---------- Marker: CHƯƠNG ngoài (IN HOA) ----------
            m = RE_CHUONG_OUT.match(text)
            if m:
                flush()
                commit_pending()
                st["chuong"] = f"Chương {m.group(1)}"
                st.update(chuong_con="", muc="", trich="", dieu="", ten_dieu="")
                rest = m.group(2).strip()
                pending_title, pending_kind = (rest, "chuong") if rest else ("", "")
                continue

            # ---------- Marker: chương CON trong một quy chế (chữ thường) ----------
            m = RE_CHUONG_IN.match(text)
            if m:
                if st["muc"]:
                    flush()
                    commit_pending()
                    st["chuong_con"] = f"Chương {m.group(1)}"
                    # Số Điều tiếp tục trong cùng quy chế → KHÔNG reset dieu
                    rest = m.group(2).strip()
                    pending_title, pending_kind = (rest, "chuong_con") if rest else ("", "")
                    continue
                # Chưa có Mục nào → coi là chương ngoài
                flush()
                commit_pending()
                st["chuong"] = f"Chương {m.group(1)}"
                st.update(chuong_con="", muc="", trich="", dieu="", ten_dieu="")
                rest = m.group(2).strip()
                pending_title, pending_kind = (rest, "chuong") if rest else ("", "")
                continue

            # ---------- Marker: MỤC A/B/C/D (tên quy chế/quy định con) ----------
            m = RE_MUC.match(text)
            if m:
                flush()
                commit_pending()
                st["muc"] = text
                st.update(trich="", dieu="", ten_dieu="", chuong_con="")
                continue

            # ---------- Marker: TRÍCH văn bản pháp quy ----------
            m = RE_TRICH.match(text)
            if m:
                flush()
                commit_pending()
                st["trich"] = text
                if not st["muc"]:
                    st["muc"] = text
                continue

            # ---------- Gộp tiêu đề đa dòng cho marker đang chờ ----------
            if pending_kind:
                if len(text) < 150 and not RE_DIEU.match(text):
                    pending_title = f"{pending_title} {text}".strip() if pending_title else text
                    continue
                commit_pending()

            append_content(text)

        elif el.tag == qn("w:tbl"):
            tbl_text = _table_text(el)
            if tbl_text:
                append_content("[BẢNG]")
                append_content(tbl_text)

    commit_pending()
    flush()

    # Nội suy số trang: tin app.xml chỉ khi số ký tự/trang hợp lý,
    # nếu không ước lượng theo FALLBACK_CHARS_PER_PAGE
    pages = result.tong_so_trang
    if pages and not (MIN_CHARS_PER_PAGE <= char_count / pages <= MAX_CHARS_PER_PAGE):
        pages = 0
    result.tong_so_trang = pages if pages else max(1, char_count // FALLBACK_CHARS_PER_PAGE)
    total = max(char_count, 1)
    for s in result.sections:
        if pages:
            s.so_trang = max(1, min(pages, s._offset * pages // total + 1))
        else:
            s.so_trang = max(1, s._offset // FALLBACK_CHARS_PER_PAGE + 1)

    result.total_chars = char_count
    result.n_sections = len(result.sections)
    return result


if __name__ == "__main__":
    import sys

    # Windows console mặc định cp1252 không in được tiếng Việt
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass

    target = sys.argv[1] if len(sys.argv) > 1 else r"data/raw/SO TAY SINH VIEN 2024-2025.docx"
    parsed = load_docx(target)
    print(f"File: {parsed.file_path}")
    print(f"Tong ky tu: {parsed.total_chars} | Trang (app.xml): {parsed.tong_so_trang}")
    print(f"So section: {parsed.n_sections}")
    print()

    print("=== Cau truc phan/chuong/muc ===")
    seen = set()
    for s in parsed.sections:
        key = (s.phan, s.chuong, s.muc)
        if key not in seen and any(key):
            seen.add(key)
            print(f"  {s.phan or '(trong)'} | {s.chuong or '-'} | {s.muc[:70] or '-'}")
    print()

    n_dieu = sum(1 for s in parsed.sections if s.dieu)
    n_bang = sum(1 for s in parsed.sections if "[BẢNG]" in s.text)
    print(f"Section co Dieu: {n_dieu} | Section chua bang: {n_bang}")
    print()

    print("=== 8 section mau ===")
    for s in parsed.sections[5:13]:
        meta = " | ".join(x for x in (s.phan, s.chuong, s.chuong_con, s.muc[:45], s.trich[:30], s.dieu) if x)
        print(f"[{meta} | tr.~{s.so_trang}]")
        print(f"  {s.text[:110]}...")
